# Llama_Syn/services/file_handler_service.py
# UPDATED FILE - Added write_file_content method

import logging
import os
from typing import Tuple, Optional

# --- Dependency Imports ---
try:
    import PyPDF2

    PYPDF2_AVAILABLE = True
except ImportError:
    PyPDF2 = None;
    PYPDF2_AVAILABLE = False
    logging.warning("FileHandlerService: PyPDF2 library not found. Install: pip install pypdf2")

try:
    import docx  # From python-docx library

    DOCX_AVAILABLE = True
except ImportError:
    docx = None;
    DOCX_AVAILABLE = False;
    logging.warning("FileHandlerService: python-docx library not found. Install: pip install python-docx")

# --- Local Imports ---
from utils import constants

logger = logging.getLogger(__name__)


class FileHandlerService:
    """Provides methods to read and extract text content from various file types."""

    def __init__(self):
        logger.info("FileHandlerService initialized.")
        if not PYPDF2_AVAILABLE:
            logger.warning("PDF handling will be unavailable.")
        if not DOCX_AVAILABLE:
            logger.warning("DOCX handling will be unavailable.")

    def read_file_content(self, file_path: str) -> Tuple[Optional[str], str, Optional[str]]:
        """
        Safely reads file content, identifies type (text, binary, error), handles PDF and DOCX.

        Args:
            file_path: The absolute path to the file.

        Returns:
            A tuple containing:
            - content (Optional[str]): The extracted text content, or None on error/binary.
            - file_type (str): "text", "binary", or "error".
            - error_msg (Optional[str]): An error message if file_type is "error", otherwise None.
        """
        display_name = os.path.basename(file_path)
        file_ext_lower = os.path.splitext(file_path)[1].lower()

        # --- File Size Check ---
        try:
            file_size = os.path.getsize(file_path)
            max_size_mb = getattr(constants, 'RAG_MAX_FILE_SIZE_MB', 50)
            max_size_bytes = max_size_mb * 1024 * 1024
            if file_size > max_size_bytes:
                err_msg = f"File > {max_size_mb}MB"
                logger.warning(f"{err_msg}: '{display_name}' ({file_size / (1024 * 1024):.2f}MB)")
                return None, "error", err_msg
            # Allow reading empty files, but chunking service might skip them.
            # if file_size == 0:
            #     logger.warning(f"File empty: '{display_name}'");
            #     return "", "text", None # Return empty string for empty text files
        except OSError as e:
            err_msg = f"OS Error (size): {e}";
            logger.error(f"Error getting size for '{display_name}': {e}");
            return None, "error", err_msg

        # --- PDF Handling ---
        if file_ext_lower == '.pdf':
            if not PYPDF2_AVAILABLE:
                logger.error(f"Cannot read PDF '{display_name}': PyPDF2 not installed.")
                return None, "error", "PyPDF2 not installed"
            try:
                content_list = []
                with open(file_path, 'rb') as pdf_file:
                    reader = PyPDF2.PdfReader(pdf_file)
                    num_pages = len(reader.pages)
                    logger.info(f"Reading {num_pages} pages from PDF: '{display_name}'")
                    for page_num in range(num_pages):
                        try:
                            page_text = reader.pages[page_num].extract_text()
                            if page_text:  # Only append if text was extracted
                                content_list.append(page_text)
                        except Exception as e_page:
                            logger.warning(
                                f"Error extracting text from page {page_num + 1} of '{display_name}': {e_page}")
                pdf_content = "\n\n--- Page Break ---\n\n".join(content_list).strip()
                if not pdf_content and num_pages > 0:  # If pages exist but no text extracted
                    logger.warning(
                        f"No text extracted from PDF: '{display_name}' (it may be image-based or protected).")
                    return None, "error", "No text extracted (image-based or protected?)"
                elif not pdf_content and num_pages == 0:  # Empty PDF
                    logger.warning(f"PDF file is empty (0 pages): '{display_name}'")
                    return "", "text", None  # Treat as empty text file
                logger.info(f"Successfully extracted text from PDF: '{display_name}' (Length: {len(pdf_content)})")
                return pdf_content, "text", None
            except PyPDF2.errors.PdfReadError as e_pdf_read:  # Specific PyPDF2 read error
                logger.error(f"PyPDF2 PdfReadError for '{display_name}': {e_pdf_read}")
                return None, "error", f"Invalid PDF: {e_pdf_read}"
            except Exception as e_pdf:
                logger.exception(f"Error reading PDF '{display_name}': {e_pdf}")
                return None, "error", f"PDF Read Error: {e_pdf}"

        # --- DOCX Handling ---
        elif file_ext_lower == '.docx':
            if not DOCX_AVAILABLE:
                logger.error(f"Cannot read DOCX '{display_name}': python-docx not installed.")
                return None, "error", "python-docx not installed"
            try:
                logger.info(f"Reading DOCX file: '{display_name}'")
                document = docx.Document(file_path)
                content_list = [p.text for p in document.paragraphs if
                                p.text.strip()]  # Ensure paragraph has stripped content
                docx_content = "\n\n".join(content_list).strip()
                if not docx_content and len(document.paragraphs) > 0:
                    logger.warning(
                        f"No text extracted from DOCX: '{display_name}' (it may be empty or structured differently).")
                    # We return "" for empty text docx, error if unreadable.
                    # If it's truly empty but readable, it's not an error.
                    # If it has structure but no text (e.g. only images), it's not an error for FileHandler.
                    return "", "text", None
                elif not docx_content and len(document.paragraphs) == 0:
                    logger.warning(f"DOCX file appears empty (0 paragraphs): '{display_name}'")
                    return "", "text", None
                logger.info(f"Successfully extracted text from DOCX: '{display_name}' (Length: {len(docx_content)})")
                return docx_content, "text", None
            except Exception as e_docx:  # Catch specific docx errors if library provides them
                logger.exception(f"Error reading DOCX '{display_name}': {e_docx}")
                return None, "error", f"DOCX Read Error: {e_docx}"

        # --- Generic Text Reading ---
        else:
            try:
                with open(file_path, "r", encoding="utf-8", errors='strict') as f:
                    content = f.read()
                    if '\x00' in content[:1024]:  # Check for null bytes
                        logger.warning(f"File likely binary (null bytes found in UTF-8 read): '{display_name}'");
                        return None, "binary", None
                    return content, "text", None
            except UnicodeDecodeError:
                logger.warning(f"UTF-8 decode failed for '{display_name}'. Trying latin-1...")
                try:
                    with open(file_path, "r", encoding="latin-1") as f:
                        content = f.read()
                        if '\x00' in content[:1024]:
                            logger.warning(f"File (read as latin-1) likely binary: '{display_name}'");
                            return None, "binary", None
                        logger.info(f"Read '{display_name}' using latin-1.");
                        return content, "text", None
                except Exception as e_fallback:
                    err_msg = f"Fallback read failed: {e_fallback}"
                    logger.warning(
                        f"Failed read '{display_name}' with fallback. Treating as binary/unreadable. Error: {err_msg}")
                    return None, "binary", err_msg
            except FileNotFoundError:
                logger.error(f"File not found during read: '{display_name}'");
                return None, "error", "File not found"
            except OSError as e_os:
                err_msg = f"OS Error (read): {e_os}";
                logger.error(f"OS error reading '{display_name}': {e_os}");
                return None, "error", err_msg
            except Exception as e_gen:
                err_msg = f"Unexpected read error: {e_gen}";
                logger.exception(f"Unexpected read error '{display_name}': {e_gen}");
                return None, "error", err_msg

    # --- NEW METHOD ---
    def write_file_content(self, file_path: str, content: str) -> Tuple[bool, Optional[str]]:
        """
        Writes string content to a file. Creates parent directories if they don't exist.

        Args:
            file_path: The absolute path to the file to be written.
            content: The string content to write.

        Returns:
            A tuple: (success_flag, error_message_or_none)
        """
        display_name = os.path.basename(file_path)
        logger.info(f"Attempting to write content to file: '{file_path}' (Length: {len(content)})")

        if not isinstance(file_path, str) or not file_path.strip():
            err = "Invalid file path provided for writing."
            logger.error(err)
            return False, err

        if not isinstance(content, str):
            # While technically you could write non-str, for our purposes it should be str.
            err = f"Invalid content type for writing to '{display_name}': Expected str, got {type(content)}."
            logger.error(err)
            return False, err

        try:
            # Ensure parent directory exists
            parent_dir = os.path.dirname(file_path)
            if parent_dir:  # Only create if parent_dir is not empty (e.g. for files in root)
                os.makedirs(parent_dir, exist_ok=True)
                logger.debug(f"Ensured directory exists: {parent_dir}")

            # Write the file
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)

            logger.info(f"Successfully wrote content to '{display_name}'.")
            return True, None
        except OSError as e_os:
            err_msg = f"OS Error writing file '{display_name}': {e_os}"
            logger.error(err_msg, exc_info=True)
            return False, err_msg
        except Exception as e_gen:
            err_msg = f"Unexpected error writing file '{display_name}': {e_gen}"
            logger.error(err_msg, exc_info=True)
            return False, err_msg
    # --- END NEW METHOD ---